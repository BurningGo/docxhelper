import os
import sys

## Add the absolute path of the parent into PYTHONPATH
from shutil import copyfile

import docx

sys.path.insert(0,os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

import unittest
import docxhelper
from docx import Document
from docx.table import Table

def _extractDocumentAndFirstTable(filePath):
    document = Document(filePath)
    tables = document.tables
    table  = tables[0]
    return document, table

def _saveDocumentWithTable(filePath):
    document = Document()

class TestTableModule(unittest.TestCase):
    def test_AlignAllCellsVerticallyCenter(self):
        srcFileName = "test_AlignAllCellsVerticallyCenter_before.docx"
        dstFileName = "test_AlignAllCellsVerticallyCenter_after.docx"
        document, tableObj = _extractDocumentAndFirstTable(srcFileName)
        docxhelper.AlignAllCellsVerticallyCenter(tableObj)
        document.save(dstFileName)

    def test_SetTableWidth_TwoColumns(self):
        srcFileName = "test_SetTableWidth_TwoColumns_before.docx"
        document, tableObj = _extractDocumentAndFirstTable(srcFileName)
        docxhelper.SetTableWidth_TwoColumns(tableObj, 2, 5)
        dstFileName_1 = "test_SetTableWidth_TwoColumns_after_2_5.docx"
        document.save(dstFileName_1)

        docxhelper.SetTableWidth_TwoColumns(tableObj, 6, 3)
        dstFileName_2 = "test_SetTableWidth_TwoColumns_after_6_3.docx"
        document.save(dstFileName_2)

if __name__ == '__main__':
    unittest.main()
    print("Check output word files.")


